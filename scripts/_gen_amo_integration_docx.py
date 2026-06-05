"""Генерация Word-документа с пояснительной запиской по интеграции amoCRM.

Конвертирует docs/amo-integration.md → docs/Пояснительная записка - интеграция amoCRM.docx.
Делает базовую разметку: заголовки, абзацы, code-блоки моноширинным, таблицы.

Запуск:
    python scripts/_gen_amo_integration_docx.py
"""
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Нужна библиотека python-docx: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def set_cell_background(cell, color_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_styled_heading(doc, text: str, level: int):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x36, 0x6B)
    return h


def add_code_block(doc, code: str):
    """Добавляет код-блок моноширинным шрифтом с серой заливкой."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    # Серая заливка фона у абзаца
    p_pr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F2F2F2")
    p_pr.append(shd)


def add_md_table(doc, rows: list):
    """rows = list of cells (list of list). Первая строка = header."""
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = cell_text
            if i == 0:
                # Header — bold + light gray bg
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
                set_cell_background(cell, "E0E6F2")


def parse_md_to_docx(md_path: Path, docx_path: Path):
    text = md_path.read_text(encoding="utf-8")
    doc = Document()

    # Document defaults
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = text.splitlines()
    i = 0
    in_code = False
    code_buf = []
    pending_table = []

    def flush_table():
        nonlocal pending_table
        if pending_table:
            # Удаляем разделитель |---|---| (вторая строка обычно)
            cleaned = [r for r in pending_table if not all(c.strip() in ("", "---") or set(c.strip()) <= {"-", ":"} for c in r)]
            add_md_table(doc, cleaned)
            pending_table = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Code blocks
        if stripped.startswith("```"):
            flush_table()
            if in_code:
                add_code_block(doc, "\n".join(code_buf))
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # Tables
        if "|" in line and stripped.startswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            pending_table.append(cells)
            i += 1
            continue
        else:
            flush_table()

        # Headings
        if stripped.startswith("# "):
            add_styled_heading(doc, stripped[2:], level=0)
        elif stripped.startswith("## "):
            add_styled_heading(doc, stripped[3:], level=1)
        elif stripped.startswith("### "):
            add_styled_heading(doc, stripped[4:], level=2)
        elif stripped.startswith("#### "):
            add_styled_heading(doc, stripped[5:], level=3)
        # Bullet lists
        elif stripped.startswith("- ") or stripped.startswith("* "):
            content = stripped[2:]
            para = doc.add_paragraph(style="List Bullet")
            apply_inline_formatting(para, content)
        elif re.match(r"^\d+\.\s", stripped):
            content = re.sub(r"^\d+\.\s", "", stripped)
            para = doc.add_paragraph(style="List Number")
            apply_inline_formatting(para, content)
        elif stripped.startswith("> "):
            # Quote
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(1.0)
            run = para.add_run(stripped[2:])
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        elif stripped == "---":
            # Horizontal rule
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            run = para.add_run("─" * 60)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        elif stripped == "":
            doc.add_paragraph()
        else:
            para = doc.add_paragraph()
            apply_inline_formatting(para, stripped)

        i += 1

    flush_table()

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    print(f"Сохранил: {docx_path}")


def apply_inline_formatting(para, text: str):
    """Применяет **bold**, `code` к тексту в абзаце."""
    # Регексп: разбиваем на токены — bold (**...**), code (`...`), обычный
    pattern = r"(\*\*[^*]+\*\*|`[^`]+`)"
    parts = re.split(pattern, text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            para.add_run(part)


if __name__ == "__main__":
    root = Path(__file__).resolve().parent.parent
    md = root / "docs" / "amo-integration.md"
    out = root / "docs" / "Пояснительная записка - интеграция amoCRM.docx"
    if not md.exists():
        print(f"Не найден: {md}", file=sys.stderr)
        sys.exit(1)
    parse_md_to_docx(md, out)
